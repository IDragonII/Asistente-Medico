import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from PIL import Image, ImageTk
import csv
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import tkintermapview
import geocoder
import webbrowser 

data = {
    'sintomas': [
        'fiebre tos dolor_de_garganta', 'fiebre tos congestión_nasal', 
        'fiebre malestar_general dolor_de_cabeza', 'fiebre dolor_de_cuerpo tos_seca', 
        'fiebre estornudos fatiga', 'fiebre tos_seca perdida_gusto', 
        'fiebre tos dificultad_respirar', 'fiebre dolor_de_cabeza perdida_olfato', 
        'fiebre tos congestión_nasal', 'fiebre y fatiga', 
        'estornudos picazon congestion_nasal', 'picazon ojos estornudos', 
        'picazon piel fiebre', 'tos estornudos picazon', 
        'picazon congestion_nasal estornudos', 'tos dolor_de_garganta', 
        'congestion_nasal estornudos fiebre', 'tos ligera malestar_general', 
        'dolor_de_garganta estornudos', 'fiebre leve tos', 
        'fiebre erupción picazón', 'tos fiebre dolor_abdominal', 
        'fiebre picazón en piel', 'fatiga tos ligera', 
        'dolor_de_cabeza mareos fatiga', 'dolor_abdominal diarrea fiebre', 
        'nauseas vomitos dolor_abdominal', 'diarrea malestar_general', 
        'fiebre y malestar abdominal', 'fatiga debilidad', 
        'dolor_de_cabeza mareos visión_doble', 'confusión mareos debilidad', 
        'fatiga debilidad temblores', 'dificultad_concentracion pérdida_memoria', 
        'dolor_de_cabeza dificultad_hablar', 'fiebre escalofríos sudoración', 
        'fiebre tos húmeda', 'dolor_pecho dificultad_respirar', 
        'fiebre sarpullido', 'fatiga cansancio extremo', 
        'nauseas dolor_de_estómago', 'vómitos mareos', 
        'tos irritativa', 'fiebre muscular', 
        'dolor_en_los_huesos', 'dificultad_para_dormir', 
        'sudoración nocturna', 'fiebre y escalofríos', 
        'dolor_en_ojos', 'mareos al levantarme', 
        'fiebre persistente', 'tos persistente'
    ],
    'enfermedad': [
        'gripe', 'gripe', 'gripe', 'gripe', 'gripe', 
        'covid', 'covid', 'covid', 'covid', 'covid', 
        'alergia', 'alergia', 'alergia', 'alergia', 'alergia', 
        'resfriado', 'resfriado', 'resfriado', 'resfriado', 'resfriado', 
        'dermatitis', 'gastroenteritis', 'dermatitis', 'anemia', 'migraña', 
        'gastroenteritis', 'gastroenteritis', 'infeccion intestinal', 'apendicitis', 'anemia', 
        'migraña', 'ACV', 'parkinson', 'alzheimer', 'ACV',
        'neumonía', 'neumonía', 'neumonía', 'neumonía', 'neumonía',
        'infección urinaria', 'infección urinaria', 'gastritis', 'gastritis',
        'insuficiencia cardíaca', 'insuficiencia cardíaca', 'acidosis', 'acidosis',
        'hipertensión', 'hipertensión', 'fibromialgia', 'fibromialgia',
    ]
}

consejos = {
    'gripe': "Descansa mucho, mantente hidratado y considera un analgésico si tienes fiebre.",
    'covid': "Sigue las pautas de salud pública y consulta a un médico si los síntomas empeoran.",
    'alergia': "Evita los alérgenos conocidos y considera antihistamínicos.",
    'resfriado': "Descansa, mantente hidratado y considera un descongestionante.",
    'dermatitis': "Mantén la piel hidratada y evita irritantes.",
    'gastroenteritis': "Hidrátate y evita alimentos grasos o picantes.",
    'anemia': "Aumenta la ingesta de hierro y consulta a un médico.",
    'migraña': "Descansa en un ambiente oscuro y tranquilo, considera analgésicos.",
    'ACV': "Busca atención médica de inmediato, es una emergencia.",
    'neumonía': "Consulta a un médico de inmediato, puede ser grave.",
    'infección urinaria': "Bebe mucha agua y consulta a un médico para antibióticos.",
    'gastritis': "Evita alimentos irritantes y considera consultar a un médico.",
    'insuficiencia cardíaca': "Consulta a un médico de inmediato para manejo adecuado.",
    'acidosis': "Busca atención médica inmediata, es una condición seria.",
    'hipertensión': "Mantén una dieta baja en sodio y consulta regularmente a un médico.",
    'fibromialgia': "Considera terapia física y manejo del estrés.",
    'artritis': "Consulta a un reumatólogo para un tratamiento adecuado."
}

df = pd.DataFrame(data)

df['sintomas'] = df['sintomas'].apply(lambda x: x.split())
X = df['sintomas']
y = df['enfermedad']

sintomas_unicos = sorted(set(sum(X, [])))

X_numerico = []
for sintomas in X:
    vector = [1 if symptom in sintomas else 0 for symptom in sintomas_unicos]
    X_numerico.append(vector)

X_train, X_test, y_train, y_test = train_test_split(X_numerico, y, test_size=0.2, random_state=42)

model = RandomForestClassifier(random_state=42)
model.fit(X_train, y_train)

def predecir_enfermedad(sintomas_ingresados):
    sintomas_vector = [1 if symptom in sintomas_ingresados else 0 for symptom in sintomas_unicos]
    probabilidades = model.predict_proba([sintomas_vector])
    return probabilidades[0]

def mostrar_prediccion():
    sintomas_usuario = sintomas_lista.get(0, tk.END)
    if not sintomas_usuario:
        resultado_text.config(state=tk.NORMAL)
        resultado_text.delete(1.0, tk.END)
        resultado_text.insert(tk.END, "Por favor, ingrese síntomas para realizar la predicción.\n")
        resultado_text.config(state=tk.DISABLED)
        return

    sintomas_usuario = list(sintomas_usuario)
    probabilidades = predecir_enfermedad(sintomas_usuario)

    resultado_text.config(state=tk.NORMAL)
    resultado_text.delete(1.0, tk.END)

    resultados = list(zip(model.classes_, probabilidades))

    resultados_ordenados = sorted(resultados, key=lambda x: x[1], reverse=True)

    for enfermedad, probabilidad in resultados_ordenados:
        resultado_text.insert(tk.END, f"{enfermedad}: {probabilidad * 100:.2f}%\n")

    resultado_text.config(state=tk.DISABLED)

def mostrar_mapa():
    g = geocoder.ip('me')
    latitud = g.latlng[0] if g.latlng else None
    longitud = g.latlng[1] if g.latlng else None

    if latitud and longitud:
        url_mapa = f"https://www.google.com/maps?q={latitud},{longitud}"
        webbrowser.open(url_mapa)
    else:
        print("No se pudo obtener la ubicación.")

def mostrar_consejos():
    sintomas_usuario = sintomas_lista.get(0, tk.END)
    if not sintomas_usuario:
        resultado_text.config(state=tk.NORMAL)
        resultado_text.delete(1.0, tk.END)
        resultado_text.insert(tk.END, "Por favor, ingrese síntomas para ver consejos.\n")
        resultado_text.config(state=tk.DISABLED)
        return

    sintomas_usuario = list(sintomas_usuario)
    probabilidades = predecir_enfermedad(sintomas_usuario)
    

    resultado_text.config(state=tk.NORMAL)
    resultado_text.delete(1.0, tk.END)

    enfermedad_max = model.classes_[probabilidades.argmax()]

    resultado_text.insert(tk.END, f"Consejos para {enfermedad_max}: {consejos.get(enfermedad_max, 'No hay consejos disponibles.')}\n")
    resultado_text.config(state=tk.DISABLED)

def seleccionar_sintoma(event=None):
    sintoma = entrada_sintomas.get().strip()
    if sintoma in sintomas_unicos and sintoma not in sintomas_lista.get(0, tk.END):
        sintomas_lista.insert(tk.END, sintoma)
        entrada_sintomas.delete(0, tk.END)

def eliminar_ultimo_sintoma():
    if sintomas_lista.size() > 0:
        sintomas_lista.delete(tk.END)

resultados_guardados = []
consejo_guardado = ""

def guardar_datos():
    global resultados_guardados, consejo_guardado
    sintomas_usuario = sintomas_lista.get(0, tk.END)
    if not sintomas_usuario:
        resultado_text.config(state=tk.NORMAL)
        resultado_text.delete(1.0, tk.END)
        resultado_text.insert(tk.END, "No hay síntomas para guardar.\n")
        resultado_text.config(state=tk.DISABLED)
        return

    probabilidades = predecir_enfermedad(list(sintomas_usuario))
    resultados_guardados = list(zip(model.classes_, probabilidades))
    enfermedad_max = model.classes_[probabilidades.argmax()]
    consejo_guardado = f"Consejos para {enfermedad_max}: {consejos.get(enfermedad_max, 'No hay consejos disponibles.')}"

    archivo_guardar = filedialog.asksaveasfilename(defaultextension=".csv", 
                                                    filetypes=[("Archivos CSV", "*.csv"),
                                                               ("Archivos PDF", "*.pdf"),
                                                               ("Archivos Word", "*.docx"),
                                                               ("Todos los archivos", "*.*")])
    if archivo_guardar:
        if archivo_guardar.endswith('.csv'):
            with open(archivo_guardar, mode='w', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                writer.writerow(["Síntomas", "Enfermedad", "Probabilidad", "Consejos"])
                for (enfermedad, probabilidad) in resultados_guardados:
                    writer.writerow(["; ".join(sintomas_usuario), enfermedad, f"{probabilidad * 100:.2f}%", consejo_guardado])  # Escribir cada síntoma, enfermedad y probabilidad
            resultado_text.config(state=tk.NORMAL)
            resultado_text.delete(1.0, tk.END)
            resultado_text.insert(tk.END, f"Datos guardados en {archivo_guardar}\n")
            resultado_text.config(state=tk.DISABLED)

        elif archivo_guardar.endswith('.pdf'):
            c = canvas.Canvas(archivo_guardar, pagesize=letter)
            c.drawString(100, 750, "Síntomas Ingresados:")
            y = 730
            for sintoma in sintomas_usuario:
                c.drawString(100, y, sintoma)
                y -= 20
            
            c.drawString(100, y, "Resultados:")
            y -= 20
            for enfermedad, probabilidad in resultados_guardados:
                c.drawString(100, y, f"{enfermedad}: {probabilidad * 100:.2f}%")
                y -= 20
            
            c.drawString(100, y, "Consejo:")
            y -= 20
            c.drawString(100, y, consejo_guardado)
            c.save()
            resultado_text.config(state=tk.NORMAL)
            resultado_text.delete(1.0, tk.END)
            resultado_text.insert(tk.END, f"Datos guardados en {archivo_guardar}\n")
            resultado_text.config(state=tk.DISABLED)

        elif archivo_guardar.endswith('.docx'):
            doc = Document()
            doc.add_heading('Síntomas Ingresados', level=1)
            for sintoma in sintomas_usuario:
                doc.add_paragraph(sintoma)
            doc.add_heading('Resultados', level=2)
            for enfermedad, probabilidad in resultados_guardados:
                doc.add_paragraph(f"{enfermedad}: {probabilidad * 100:.2f}%")
            doc.add_heading('Consejo', level=2)
            doc.add_paragraph(consejo_guardado)
            doc.save(archivo_guardar)
            resultado_text.config(state=tk.NORMAL)
            resultado_text.delete(1.0, tk.END)
            resultado_text.insert(tk.END, f"Datos guardados en {archivo_guardar}\n")
            resultado_text.config(state=tk.DISABLED)


ventana = tk.Tk()
ventana.title("Asistente de Salud")
ventana.geometry("800x600")
ventana.configure(bg="#f0f0f0")

menu_bar = tk.Menu(ventana)
ventana.config(menu=menu_bar)

menu_archivo = tk.Menu(menu_bar, tearoff=0)
menu_archivo.add_command(label="Nuevo", command=lambda: print("Nuevo seleccionado"))
menu_archivo.add_command(label="Abrir", command=lambda: print("Abrir seleccionado"))
menu_archivo.add_command(label="Guardar", command=guardar_datos)
menu_archivo.add_separator()
menu_archivo.add_command(label="Salir", command=ventana.quit)
menu_bar.add_cascade(label="Archivo", menu=menu_archivo)

menu_hospitales = tk.Menu(menu_bar, tearoff=0)
menu_hospitales.add_command(label="Mapa", command=mostrar_mapa)
menu_bar.add_cascade(label="Hospitales", menu=menu_hospitales)

menu_ayuda = tk.Menu(menu_bar, tearoff=0)
menu_ayuda.add_command(label="Acerca de", command=lambda: print("Acerca de seleccionado"))
menu_bar.add_cascade(label="Ayuda", menu=menu_ayuda)

inicio_frame = tk.Frame(ventana, bg="#f0f0f0")
inicio_frame.pack(fill="both", expand=True)

try:
    imagen_inicio = Image.open("salud_inicio.png")
    imagen_inicio = imagen_inicio.resize((200, 200), Image.ANTIALIAS)
    imagen_inicio_tk = ImageTk.PhotoImage(imagen_inicio)
    label_imagen = tk.Label(inicio_frame, image=imagen_inicio_tk, bg="#f0f0f0")
    label_imagen.pack(pady=20)
except Exception as e:
    print(f"No se pudo cargar la imagen: {e}")

mensaje_inicio = tk.Label(inicio_frame, text="¡Bienvenido al Asistente de Salud!", font=("Helvetica", 18, "bold"), bg="#f0f0f0", fg="#333")
mensaje_inicio.pack(pady=20)

descripcion = tk.Label(inicio_frame, text="Proporcione sus síntomas y reciba predicciones de salud y consejos.", 
                       font=("Helvetica", 12), bg="#f0f0f0", fg="#666")
descripcion.pack(pady=10)

boton_iniciar = tk.Button(inicio_frame, text="Iniciar", font=("Helvetica", 14), command=lambda: [inicio_frame.destroy(), crear_interfaz_principal()], bg="#4CAF50", fg="white", width=10, height=2)
boton_iniciar.pack(pady=20)

def crear_interfaz_principal():

    frame = tk.Frame(ventana, bg="#f0f0f0")
    frame.pack(pady=10, padx=10, fill="both", expand=True)


    frame_sintomas = tk.LabelFrame(frame, text="Ingrese sus síntomas", font=("Helvetica", 14, "bold"), bg="#f0f0f0", fg="#333", padx=20, pady=10)
    frame_sintomas.pack(side="left", fill="both", expand=True, padx=10, pady=10)

    global entrada_sintomas
    entrada_sintomas = tk.Entry(frame_sintomas, font=("Helvetica", 12), width=30, borderwidth=2, relief="groove")
    entrada_sintomas.pack(pady=10)

    entrada_sintomas.bind("<Return>", seleccionar_sintoma)

    global sintomas_lista
    sintomas_lista = tk.Listbox(frame_sintomas, font=("Helvetica", 12), width=30, height=10, borderwidth=2, relief="groove")
    sintomas_lista.pack(pady=10)

    boton_agregar = tk.Button(frame_sintomas, text="Agregar síntoma", font=("Helvetica", 12), command=lambda: seleccionar_sintoma(None), bg="#4CAF50", fg="white")
    boton_agregar.pack(pady=5)
    boton_eliminar = tk.Button(frame_sintomas, text="Eliminar síntoma", font=("Helvetica", 12), command=eliminar_ultimo_sintoma, bg="#f44336", fg="white")
    boton_eliminar.pack(pady=5)

    frame_resultados = tk.LabelFrame(frame, text="Resultados", font=("Helvetica", 14, "bold"), bg="#f0f0f0", fg="#333", padx=20, pady=10)
    frame_resultados.pack(side="right", fill="both", expand=True, padx=10, pady=10)

    global resultado_text
    resultado_text = tk.Text(frame_resultados, font=("Helvetica", 12), width=40, height=15, state=tk.DISABLED, bg="#ffffff", borderwidth=2, relief="groove")
    resultado_text.pack(pady=10)

    boton_prediccion = tk.Button(frame_resultados, text="Mostrar predicción", font=("Helvetica", 12), command=mostrar_prediccion, bg="#2196F3", fg="white")
    boton_prediccion.pack(pady=5)
    boton_consejos = tk.Button(frame_resultados, text="Mostrar consejos", font=("Helvetica", 12), command=mostrar_consejos, bg="#FF9800", fg="white")
    boton_consejos.pack(pady=5)

ventana.mainloop()
