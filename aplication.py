import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from PIL import Image, ImageTk
import csv
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import tkintermapview
import geocoder
import webbrowser 

data = {
    'sintomas': [
        'fiebre tos dolor_de_garganta', 'fiebre tos congestión_nasal',
        'fiebre malestar_general dolor_de_cabeza', 'fiebre dolor_de_cuerpo tos_seca',
        'fiebre estornudos fatiga', 'fiebre tos_seca pérdida_de_gusto',
        'fiebre tos dificultad_respirar', 'fiebre dolor_de_cabeza pérdida_de_olfato',
        'fiebre tos congestión_nasal', 'fiebre y fatiga',
        'estornudos picazón congestión_nasal', 'picazón ojos estornudos',
        'picazón piel fiebre', 'tos estornudos picazón',
        'picazón congestión_nasal estornudos', 'tos dolor_de_garganta',
        'congestión_nasal estornudos fiebre', 'tos ligera malestar_general',
        'dolor_de_garganta estornudos', 'fiebre leve tos',
        'fiebre erupción picazón', 'tos fiebre dolor_abdominal',
        'fiebre picazón en piel', 'fatiga tos ligera',
        'dolor_de_cabeza mareos fatiga', 'dolor_abdominal diarrea fiebre',
        'nauseas vómitos dolor_abdominal', 'diarrea malestar_general',
        'fiebre y malestar abdominal', 'fatiga debilidad',
        'dolor_de_cabeza mareos visión_doble', 'confusión mareos debilidad',
        'fatiga debilidad temblores', 'dificultad_concentracion pérdida_memoria',
        'dolor_de_cabeza dificultad_hablar', 'fiebre escalofríos sudoración',
        'fiebre tos húmeda', 'dolor_pecho dificultad_respirar',
        'fiebre sarpullido', 'fatiga cansancio extremo',
        'nauseas dolor_de_estómago', 'vómitos mareos',
        'tos irritativa', 'fiebre muscular',
        'dolor_en_los_huesos', 'dificultad_para_dormir',
        'sudoración nocturna', 'fiebre y escalofríos',
        'dolor_en_ojos', 'mareos al levantarme',
        'fiebre persistente', 'tos persistente', 
        'dolor_articular rigidez', 'náuseas fatiga debilidad',
        'erupción cutánea dolor_de_cabeza', 'fiebre tos secreción_nasal',
        'pérdida_de_apetito náuseas', 'diarrea vómitos deshidratación',
        'dificultad_para_concentrarse fatiga', 'dolor_en_los_hombros dolor_de_cabeza',
        'tos_cronica dificultad_respirar', 'fiebre cansancio extremo', 
        'vomitos dolor_en_estómago', 'picazón fiebre',
        'mareos confusión dificultad_para_hablar', 'sudores_nocturnos fatiga',
        'dificultad_para_dormir cansancio', 'tos seca fiebre',
        'dolor_pecho presión', 'palpitaciones mareos',
    ],
    'enfermedad': [
        'gripe', 'gripe', 'gripe', 'gripe', 'gripe',
        'covid-19', 'covid-19', 'covid-19', 'covid-19', 'covid-19',
        'alergia', 'alergia', 'alergia', 'alergia', 'alergia',
        'resfriado', 'resfriado', 'resfriado', 'resfriado', 'resfriado',
        'dermatitis', 'gastroenteritis', 'dermatitis', 'anemia', 'migraña',
        'gastroenteritis', 'gastroenteritis', 'infección intestinal', 'apendicitis', 'anemia',
        'migraña', 'ACV', 'Parkinson', 'Alzheimer', 'ACV',
        'neumonía', 'neumonía', 'neumonía', 'neumonía', 'neumonía',

        'infección urinaria', 'infección urinaria', 'gastritis', 'gastritis',
        'insuficiencia cardíaca', 'insuficiencia cardíaca', 'acidosis', 'acidosis',
        'hipertensión', 'hipertensión', 'fibromialgia', 'fibromialgia',
        'artritis', 'insuficiencia renal', 'hepatitis', 'faringitis',
        'cistitis', 'insuficiencia respiratoria', 'trastornos de ansiedad', 'depresión',
        'esclerosis múltiple', 'asma', 'bronquitis crónica', 'neumonía atípica',
        'síndrome de fatiga crónica', 'hipotiroidismo', 'hipertiroidismo',
        'síndrome premenstrual', 'fibrosis pulmonar', 'cáncer de pulmón',
    ]
}


consejos = {
    'gripe': "Descansa mucho, mantente hidratado y considera un analgésico si tienes fiebre. Si los síntomas persisten más de una semana, consulta a un médico.",
    'covid-19': "Sigue las pautas de salud pública, realiza pruebas si es necesario y consulta a un médico si los síntomas empeoran.",
    'alergia': "Evita los alérgenos conocidos, considera antihistamínicos y consulta a un médico si los síntomas son severos o persistentes.",
    'resfriado': "Descansa, mantente hidratado y considera un descongestionante. Si los síntomas no mejoran en 10 días, consulta a un médico.",
    'dermatitis': "Mantén la piel hidratada y evita irritantes. Consulta a un dermatólogo si la irritación persiste.",
    'gastroenteritis': "Hidrátate con soluciones electrolíticas y evita alimentos grasos o picantes. Si los síntomas son graves, consulta a un médico.",
    'anemia': "Aumenta la ingesta de hierro con alimentos como espinacas y carne roja, y consulta a un médico para un análisis de sangre.",
    'migraña': "Descansa en un ambiente oscuro y tranquilo, y considera analgésicos. Si las migrañas son frecuentes, consulta a un médico.",
    'ACV': "Busca atención médica de inmediato, es una emergencia. Recuerda los signos: debilidad en un lado del cuerpo, confusión, dificultad para hablar.",
    'neumonía': "Consulta a un médico de inmediato; puede ser grave. La atención temprana es clave.",
    'infección urinaria': "Bebe mucha agua y consulta a un médico para antibióticos. No ignores los síntomas como ardor al orinar.",
    'gastritis': "Evita alimentos irritantes y considera consultar a un médico si los síntomas persisten o empeoran.",
    'insuficiencia cardíaca': "Consulta a un médico de inmediato para manejo adecuado; los síntomas pueden incluir dificultad para respirar y fatiga extrema.",
    'acidosis': "Busca atención médica inmediata, ya que es una condición seria que requiere tratamiento especializado.",
    'hipertensión': "Mantén una dieta baja en sodio, haz ejercicio regularmente y consulta a un médico para chequeos regulares.",
    'fibromialgia': "Considera terapia física y manejo del estrés, y consulta a un médico para un plan de tratamiento adecuado.",
    'artritis': "Consulta a un reumatólogo para un tratamiento adecuado y maneja los síntomas con ejercicio y medicamentos antiinflamatorios.",
    'insuficiencia renal': "Mantén un control de la dieta y consulta a un nefrólogo; los síntomas pueden incluir hinchazón y fatiga.",
    'hepatitis': "Evita el alcohol, mantén una dieta equilibrada y consulta a un médico para seguimiento y tratamiento.",
    'faringitis': "Garganta con agua salada y consulta a un médico si los síntomas persisten más de 3 días.",
    'cistitis': "Bebe mucha agua y consulta a un médico para antibióticos; los síntomas pueden incluir dolor al orinar.",
    'insuficiencia respiratoria': "Busca atención médica de inmediato; esto puede ser una emergencia médica.",
    'trastornos de ansiedad': "Considera terapia y técnicas de relajación; consulta a un profesional de la salud mental si los síntomas interfieren con tu vida diaria.",
    'depresión': "Consulta a un profesional de la salud mental; no ignores los síntomas como tristeza persistente y pérdida de interés en actividades.",
    'esclerosis múltiple': "Consulta a un neurólogo para un manejo adecuado; los síntomas pueden incluir debilidad y problemas de visión.",
    'asma': "Usa medicamentos de rescate y evita desencadenantes; consulta a un médico para un plan de manejo.",
    'bronquitis crónica': "Consulta a un médico para tratamiento y manejo de síntomas; evitar el tabaco es esencial.",
    'neumonía atípica': "Consulta a un médico de inmediato para tratamiento; no ignores síntomas como tos persistente y fiebre.",
    'síndrome de fatiga crónica': "Consulta a un médico para manejo de síntomas y descartar otras condiciones.",
    'hipotiroidismo': "Consulta a un endocrinólogo para tratamiento; los síntomas incluyen fatiga y aumento de peso.",
    'hipertiroidismo': "Busca atención médica para manejo adecuado; los síntomas pueden incluir pérdida de peso inexplicada y ansiedad.",
    'síndrome premenstrual': "Considera ejercicio regular y técnicas de manejo del estrés; consulta a un médico si los síntomas son severos.",
    'fibrosis pulmonar': "Consulta a un neumólogo para tratamiento adecuado; los síntomas pueden incluir dificultad para respirar.",
    'cáncer de pulmón': "Consulta a un oncólogo para un plan de tratamiento; los síntomas pueden incluir tos persistente y pérdida de peso.",
    'enfermedad de Crohn': "Consulta a un gastroenterólogo para manejo de síntomas; los síntomas pueden incluir dolor abdominal y diarrea crónica.",
    'colitis ulcerosa': "Consulta a un gastroenterólogo para tratamiento; los síntomas incluyen diarrea con sangre y dolor abdominal.",
    'síndrome metabólico': "Mantén un estilo de vida saludable y consulta a un médico; esto puede aumentar el riesgo de enfermedades cardiovasculares."
}


df = pd.DataFrame(data)

df['sintomas'] = df['sintomas'].apply(lambda x: x.split())
X = df['sintomas']
y = df['enfermedad']

sintomas_unicos = sorted(set(sum(X, [])))

X_numerico = []
for sintomas in X:
    vector = [1 if symptom in sintomas else 0 for symptom in sintomas_unicos]
    X_numerico.append(vector)

X_train, X_test, y_train, y_test = train_test_split(X_numerico, y, test_size=0.2, random_state=42)

model = RandomForestClassifier(random_state=42)
model.fit(X_train, y_train)

def predecir_enfermedad(sintomas_ingresados):
    sintomas_vector = [1 if symptom in sintomas_ingresados else 0 for symptom in sintomas_unicos]
    probabilidades = model.predict_proba([sintomas_vector])
    return probabilidades[0]

def mostrar_prediccion():
    sintomas_usuario = sintomas_lista.get(0, tk.END)
    if not sintomas_usuario:
        messagebox.showinfo("Advertencia", "Por favor, ingrese síntomas para realizar la predicción.")
        return

    sintomas_usuario = list(sintomas_usuario)
    probabilidades = predecir_enfermedad(sintomas_usuario)

    resultado_text.config(state=tk.NORMAL)
    resultado_text.delete(1.0, tk.END)

    resultados = list(zip(model.classes_, probabilidades))

    resultados_ordenados = sorted(resultados, key=lambda x: x[1], reverse=True)

    for enfermedad, probabilidad in resultados_ordenados:
        resultado_text.insert(tk.END, f"{enfermedad}: {probabilidad * 100:.2f}%\n")

    resultado_text.config(state=tk.DISABLED)

def mostrar_mapa():
    g = geocoder.ip('me')
    latitud = g.latlng[0] if g.latlng else None
    longitud = g.latlng[1] if g.latlng else None

    if latitud and longitud:
        url_mapa = f"https://www.google.com/maps?q={latitud},{longitud}"
        webbrowser.open(url_mapa)
    else:
        print("No se pudo obtener la ubicación.")

def centrar_ventana(ventana):
    ventana.update_idletasks()
    width = ventana.winfo_width()
    height = ventana.winfo_height()
    x = (ventana.winfo_screenwidth() // 2) - (width // 2)
    y = (ventana.winfo_screenheight() // 2) - (height // 2)
    ventana.geometry(f'{width}x{height}+{x}+{y}')

def mostrar_consejos():
    sintomas_usuario = sintomas_lista.get(0, tk.END)
    if not sintomas_usuario:
        messagebox.showinfo("Advertencia", "Por favor, ingrese síntomas para ver consejos.")
        return

    sintomas_usuario = [s.replace("_", " ") for s in list(sintomas_usuario)]
    probabilidades = predecir_enfermedad(sintomas_usuario)

    resultado_text.config(state=tk.NORMAL)
    resultado_text.delete(1.0, tk.END)

    enfermedad_max = model.classes_[probabilidades.argmax()]

    resultado_text.insert(tk.END, f"Consejos para {enfermedad_max}: {consejos.get(enfermedad_max, 'No hay consejos disponibles.')}\n")
    resultado_text.config(state=tk.DISABLED)

def seleccionar_sintoma(event=None):
    sintoma = entrada_sintomas.get().strip()
    sintoma_modificado = sintoma.replace(" ", "_")
    if sintoma_modificado in sintomas_unicos and sintoma_modificado not in sintomas_lista.get(0, tk.END):
        sintomas_lista.insert(tk.END, sintoma)  # Muestra el sintoma con espacio
        entrada_sintomas.delete(0, tk.END)

def eliminar_ultimo_sintoma():
    if sintomas_lista.size() > 0:
        sintomas_lista.delete(tk.END)

resultados_guardados = []
consejo_guardado = ""

def guardar_datos():
    global resultados_guardados, consejo_guardado
    sintomas_usuario = sintomas_lista.get(0, tk.END)
    if not sintomas_usuario:
        resultado_text.config(state=tk.NORMAL)
        resultado_text.delete(1.0, tk.END)
        resultado_text.insert(tk.END, "No hay síntomas para guardar.\n")
        resultado_text.config(state=tk.DISABLED)
        return

    probabilidades = predecir_enfermedad(list(sintomas_usuario))
    resultados_guardados = list(zip(model.classes_, probabilidades))
    enfermedad_max = model.classes_[probabilidades.argmax()]
    consejo_guardado = f"Consejos para {enfermedad_max}: {consejos.get(enfermedad_max, 'No hay consejos disponibles.')}"

    archivo_guardar = filedialog.asksaveasfilename(defaultextension=".csv", 
                                                    filetypes=[("Archivos CSV", "*.csv"),
                                                               ("Archivos PDF", "*.pdf"),
                                                               ("Archivos Word", "*.docx"),
                                                               ("Todos los archivos", "*.*")])
    if archivo_guardar:
        if archivo_guardar.endswith('.csv'):
            with open(archivo_guardar, mode='w', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                writer.writerow(["Síntomas", "Enfermedad", "Probabilidad", "Consejos"])
                for (enfermedad, probabilidad) in resultados_guardados:
                    writer.writerow(["; ".join(sintomas_usuario), enfermedad, f"{probabilidad * 100:.2f}%", consejo_guardado])  # Escribir cada síntoma, enfermedad y probabilidad
            resultado_text.config(state=tk.NORMAL)
            resultado_text.delete(1.0, tk.END)
            resultado_text.insert(tk.END, f"Datos guardados en {archivo_guardar}\n")
            resultado_text.config(state=tk.DISABLED)

        elif archivo_guardar.endswith('.pdf'):
            c = canvas.Canvas(archivo_guardar, pagesize=letter)
            c.drawString(100, 750, "Síntomas Ingresados:")
            y = 730
            for sintoma in sintomas_usuario:
                c.drawString(100, y, sintoma)
                y -= 20
            
            c.drawString(100, y, "Resultados:")
            y -= 20
            for enfermedad, probabilidad in resultados_guardados:
                c.drawString(100, y, f"{enfermedad}: {probabilidad * 100:.2f}%")
                y -= 20

            c.drawString(100, y, "Consejo:")
            y -= 20
            if y < 50:
                c.showPage()
                y = 750
            
            c.drawString(100, y, consejo_guardado)
            c.save()
            resultado_text.config(state=tk.NORMAL)
            resultado_text.delete(1.0, tk.END)
            resultado_text.insert(tk.END, f"Datos guardados en {archivo_guardar}\n")
            resultado_text.config(state=tk.DISABLED)

        elif archivo_guardar.endswith('.docx'):
            doc = Document()
            doc.add_heading('Síntomas Ingresados', level=1)
            for sintoma in sintomas_usuario:
                doc.add_paragraph(sintoma)
            doc.add_heading('Resultados', level=2)
            for enfermedad, probabilidad in resultados_guardados:
                doc.add_paragraph(f"{enfermedad}: {probabilidad * 100:.2f}%")
            doc.add_heading('Consejo', level=2)
            doc.add_paragraph(consejo_guardado)
            doc.save(archivo_guardar)
            resultado_text.config(state=tk.NORMAL)
            resultado_text.delete(1.0, tk.END)
            resultado_text.insert(tk.END, f"Datos guardados en {archivo_guardar}\n")
            resultado_text.config(state=tk.DISABLED)


ventana = tk.Tk()
ventana.title("Asistente de Salud")
ventana.geometry("800x600")
ventana.configure(bg="#f0f0f0")

menu_bar = tk.Menu(ventana)
ventana.config(menu=menu_bar)

menu_archivo = tk.Menu(menu_bar, tearoff=0)
menu_archivo.add_command(label="Nuevo", command=lambda: print("Nuevo seleccionado"))
menu_archivo.add_command(label="Abrir", command=lambda: print("Abrir seleccionado"))
menu_archivo.add_command(label="Guardar", command=guardar_datos)
menu_archivo.add_separator()
menu_archivo.add_command(label="Salir", command=ventana.quit)
menu_bar.add_cascade(label="Archivo", menu=menu_archivo)

menu_hospitales = tk.Menu(menu_bar, tearoff=0)
menu_hospitales.add_command(label="Mapa", command=mostrar_mapa)
menu_bar.add_cascade(label="Hospitales", menu=menu_hospitales)

menu_ayuda = tk.Menu(menu_bar, tearoff=0)
menu_ayuda.add_command(label="Acerca de", command=lambda: print("Acerca de seleccionado"))
menu_bar.add_cascade(label="Ayuda", menu=menu_ayuda)

inicio_frame = tk.Frame(ventana, bg="#f0f0f0")
inicio_frame.pack(fill="both", expand=True)

try:
    imagen_inicio = Image.open("salud_inicio.jpg")
    imagen_inicio = imagen_inicio.resize((200, 200), Image.Resampling.LANCZOS)
    imagen_inicio_tk = ImageTk.PhotoImage(imagen_inicio)
    label_imagen = tk.Label(inicio_frame, image=imagen_inicio_tk, bg="#f0f0f0")
    label_imagen.pack(pady=20)
except Exception as e:
    print(f"No se pudo cargar la imagen: {e}")

mensaje_inicio = tk.Label(inicio_frame, text="¡Bienvenido al Asistente de Salud!", font=("Helvetica", 18, "bold"), bg="#f0f0f0", fg="#333")
mensaje_inicio.pack(pady=20)

mensaje_inicio = tk.Label(inicio_frame, text="¡Como tu salud nos importa!", font=("Helvetica", 18, "bold"), bg="#f0f0f0", fg="#333")
mensaje_inicio.pack(pady=20)

descripcion = tk.Label(inicio_frame, text="Proporcionenos sus síntomas y le diremos que enfermedad podria tener.", 
                       font=("Helvetica", 12), bg="#f0f0f0", fg="#666")
descripcion.pack(pady=10)

boton_iniciar = tk.Button(inicio_frame, text="Iniciar", font=("Helvetica", 14), command=lambda: [inicio_frame.destroy(), crear_interfaz_principal()], bg="#4CAF50", fg="white", width=10, height=2)
boton_iniciar.pack(pady=20)

def crear_interfaz_principal():

    frame = tk.Frame(ventana, bg="#f0f0f0")
    frame.pack(pady=10, padx=10, fill="both", expand=True)


    frame_sintomas = tk.LabelFrame(frame, text="¿Que sintomas tiene?", font=("Helvetica", 14, "bold"), bg="#f0f0f0", fg="#333", padx=20, pady=10)
    frame_sintomas.pack(side="left", fill="both", expand=True, padx=10, pady=10)

    global entrada_sintomas
    entrada_sintomas = tk.Entry(frame_sintomas, font=("Helvetica", 12), width=30, borderwidth=2, relief="groove")
    entrada_sintomas.pack(pady=10)

    entrada_sintomas.bind("<Return>", seleccionar_sintoma)

    global sintomas_lista
    sintomas_lista = tk.Listbox(frame_sintomas, font=("Helvetica", 12), width=30, height=10, borderwidth=2, relief="groove")
    sintomas_lista.pack(pady=10)

    boton_agregar = tk.Button(frame_sintomas, text="Agregar síntoma", font=("Helvetica", 12), command=lambda: seleccionar_sintoma(None), bg="#4CAF50", fg="white")
    boton_agregar.pack(pady=5)
    boton_eliminar = tk.Button(frame_sintomas, text="Eliminar síntoma", font=("Helvetica", 12), command=eliminar_ultimo_sintoma, bg="#f44336", fg="white")
    boton_eliminar.pack(pady=5)

    frame_resultados = tk.LabelFrame(frame, text="Resultados", font=("Helvetica", 14, "bold"), bg="#f0f0f0", fg="#333", padx=20, pady=10)
    frame_resultados.pack(side="right", fill="both", expand=True, padx=10, pady=10)

    global resultado_text
    resultado_text = tk.Text(frame_resultados, font=("Helvetica", 12), width=40, height=15, state=tk.DISABLED, bg="#ffffff", borderwidth=2, relief="groove")
    resultado_text.pack(pady=10)

    boton_prediccion = tk.Button(frame_resultados, text="Mostrar posibles enfermedades", font=("Helvetica", 12), command=mostrar_prediccion, bg="#2196F3", fg="white")
    boton_prediccion.pack(pady=5)
    boton_consejos = tk.Button(frame_resultados, text="Que me aconsejas", font=("Helvetica", 12), command=mostrar_consejos, bg="#FF9800", fg="white")
    boton_consejos.pack(pady=5)

ventana.mainloop()
